import os
import json
import logging
import asyncio
import re
from typing import Dict, Optional, List, Type
from pathlib import Path

from docx import Document
from openai import AsyncOpenAI
import tiktoken
from pydantic import BaseModel

# PDF processing libraries
try:
    import fitz  # PyMuPDF
    import pytesseract
    from PIL import Image
    import io
except ImportError:
    fitz = None
    pytesseract = None
    Image = None
    io = None

logger = logging.getLogger(__name__)

class GenericDataExtractor:
    """
    Generalized data extractor using OpenAI Structured Outputs (Async).
    Supports dynamic response schemas and configurable model parameters.
    """
    
    def __init__(
        self, 
        model_id: str = "gpt-4o-mini",
        api_key: Optional[str] = None, 
        max_tokens: int = 100000,
        token_overlap: int = 500,
        response_model: Optional[Type[BaseModel]] = None,
        scan_density_threshold: int = 100
    ):
        self.api_key = api_key
        self.model = model_id
        self.max_tokens = max_tokens
        self.token_overlap = token_overlap
        self.response_model = response_model
        self.scan_density_threshold = scan_density_threshold
        
        self.client = None
        self.tokenizer = None
        
    async def initialize(self):
        """
        Async initialization for heavy tasks (tokenizer loading, client setup).
        """
        # Initialize OpenAI Client
        self.api_key = self.api_key or os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            raise ValueError("OpenAI API key not found. Set OPENAI_API_KEY in .env")

        try:
            self.client = AsyncOpenAI(api_key=self.api_key)
        except Exception as e:
            logger.warning(f"Error initializing OpenAI client: {e}. Trying alternative initialization...")
            import httpx
            self.client = AsyncOpenAI(
                api_key=self.api_key,
                http_client=httpx.AsyncClient()
            )

        # Load Tokenizer (run in thread to avoid blocking event loop)
        try:
            logger.info("Loading o200k_base tokenizer...")
            self.tokenizer = await asyncio.to_thread(tiktoken.get_encoding, "o200k_base")
            logger.info("o200k_base tokenizer loaded successfully.")
        except Exception:
            logger.warning("Failed to load o200k_base tokenizer. Falling back to cl100k_base.")
            self.tokenizer = await asyncio.to_thread(tiktoken.get_encoding, "cl100k_base")
            logger.info("cl100k_base tokenizer loaded successfully.")
            
        if fitz is None:
            logger.warning("PyMuPDF (fitz) is not installed. PDF extraction will fail.")
        
        if pytesseract is None:
            logger.warning("Tesseract OCR (pytesseract) is not installed. Scanned PDFs cannot be processed.")

    @staticmethod
    def sanitize_data(data):
        if data is None:
            return None
        if isinstance(data, str):
            sanitized = data.replace('\x00', '')
            sanitized = re.sub(r'[\x01-\x08\x0B\x0C\x0E-\x1F\x7F]', '', sanitized)
            return sanitized
        elif isinstance(data, dict):
            return {key: GenericDataExtractor.sanitize_data(value) for key, value in data.items()}
        elif isinstance(data, list):
            return [GenericDataExtractor.sanitize_data(item) for item in data]
        return data

    def count_tokens(self, text: str) -> int:
        return len(self.tokenizer.encode(text))

    def chunk_text(self, text: str) -> List[str]:
        tokens = self.tokenizer.encode(text)
        total_tokens = len(tokens)
        
        if self.max_tokens >= total_tokens:
            return [text]
        
        chunks = []
        start_idx = 0
        while start_idx < total_tokens:
            end_idx = min(start_idx + self.max_tokens, total_tokens)
            chunks.append(self.tokenizer.decode(tokens[start_idx:end_idx]))
            if end_idx == total_tokens:
                break
            start_idx = end_idx - self.token_overlap
        return chunks

    def _extract_from_docx(self, file_path: str) -> str:
        doc = Document(file_path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)

    def _extract_from_txt(self, path: Path) -> str:
        try:
            return path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            return path.read_text(encoding='latin-1')

    def _extract_scanned_pdf(self, doc) -> str:
        if not pytesseract or not Image or not io:
            logger.warning("OCR dependencies (pytesseract, Pillow) not found. Returning empty/partial text.")
            return "" 
            
        ocr_parts = []
        for i, page in enumerate(doc):
            logger.info(f"Performing OCR on page {i+1}/{len(doc)}")
            pix = page.get_pixmap()
            img_bytes = pix.tobytes("png")
            image = Image.open(io.BytesIO(img_bytes))
            text = pytesseract.image_to_string(image, lang='eng')
            ocr_parts.append(text)
        return "\n\n".join(ocr_parts)

    def _extract_from_pdf(self, file_path: str) -> str:
        if not fitz:
            raise ImportError("PyMuPDF (fitz) is not installed.")
            
        doc = fitz.open(file_path)
        try:
            parts = []
            total_chars = 0
            page_count = len(doc)
            
            for page in doc:
                text = page.get_text()
                parts.append(text)
                total_chars += len(text.strip())
            
            avg_chars = total_chars / page_count if page_count > 0 else 0
            logger.info(f"PDF extraction stats: {page_count} pages, {total_chars} chars, avg {avg_chars:.2f} chars/page")
            
            if avg_chars < self.scan_density_threshold:
                logger.info(f"Average characters per page ({avg_chars:.2f}) is below threshold ({self.scan_density_threshold}). Assuming scanned PDF.")
                return self._extract_scanned_pdf(doc)
                
            return "\n".join(parts)
        finally:
            doc.close()

    def extract_text(self, file_path: str) -> str:
        path = Path(file_path)
        suffix = path.suffix.lower()
        
        if suffix == '.docx':
            return self._extract_from_docx(file_path)
        elif suffix == '.pdf':
            return self._extract_from_pdf(file_path)
        elif suffix == '.txt':
            return self._extract_from_txt(path)
        
        raise ValueError(f"Unsupported format: {suffix}")

    async def extract_structured(self, text: str, previous_data: Optional[Dict] = None) -> Dict:
        system_prompt = "You are a specialized data extraction agent. Extract all required information from the provided text."
        if previous_data:
            system_prompt += f"\n\nPrevious data (merge and update):\n{json.dumps(previous_data, indent=2)}"
        
        try:
            completion = await self.client.chat.completions.parse(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": text}
                ],
                response_format=self.response_model,
            )
            
            resp = completion.choices[0].message
            if resp.refusal:
                return {"status": "failed", "message": f"Refusal: {resp.refusal}"}
            
            if resp.parsed:
                return {
                    "status": "success", 
                    "data": self.sanitize_data(resp.parsed.model_dump()),
                    "message": "Extraction successful"
                }
            return {"status": "failed", "message": "Failed to parse response"}
            
        except Exception as e:
            logger.error(f"Extraction error: {e}")
            return {"status": "failed", "message": str(e)}

    async def process_file(self, file_path: str) -> Dict:
        try:
            logger.info(f"Extracting text from file: {file_path}")
            text = await asyncio.to_thread(self.extract_text, file_path)
            if not text.strip():
                logger.warning(f"No text content found in file: {file_path}")
                return {"status": "failed", "message": "No text content found"}
            
            text_length = len(text)
            total_tokens = self.count_tokens(text)
            logger.info(f"Text extracted successfully - Length: {text_length} chars, Tokens: {total_tokens}")
            
            logger.info(f"Chunking text with max_tokens={self.max_tokens}, overlap={self.token_overlap}")
            chunks = self.chunk_text(text)
            num_chunks = len(chunks)
            logger.info(f"Document split into {num_chunks} chunk(s)")
            
            accumulated = None
            for idx, chunk in enumerate(chunks, 1):
                chunk_tokens = self.count_tokens(chunk)
                logger.info(f"Processing chunk {idx}/{num_chunks} ({chunk_tokens} tokens)...")
                
                result = await self.extract_structured(chunk, accumulated)
                
                if result["status"] == "success":
                    logger.info(f"Chunk {idx}/{num_chunks} processed successfully")
                    accumulated = result["data"]
                else:
                    logger.error(f"Chunk {idx}/{num_chunks} failed: {result.get('message')}")
                    if accumulated:
                        return {"status": "failed", "data": accumulated, "message": f"Partial failure at chunk {idx}/{num_chunks}: {result['message']}"}
                    return result
            
            logger.info(f"All {num_chunks} chunks processed successfully")
            return {"status": "success", "data": accumulated, "message": f"All {num_chunks} chunks processed"}
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {e}", exc_info=True)
            return {"status": "failed", "message": str(e)}

    async def close(self):
        await self.client.close()
